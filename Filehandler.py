import glob
import re
from time import sleep, ctime
import threading
import docx
from docx import Document
import BaiduTranslator
import os


# coding=utf8

def remove_doc_space(para):
    para_text = re.sub(u'[\u3000,\xa0]', u'', para.text)  # remove em space
    para_text = para_text.replace(' ', '')  # remove en space
    para_text = para_text.replace("\n", '')  # remove empty rows
    return para_text


def remove_txt_space(para):
    if isinstance(para, str):
        para = unicode(para, "utf-8")
    para_text = re.sub(u'[\u3000,\xa0]', u'', para)  # remove em space
    para_text = para_text.replace(' ', '')  # remove en space
    para_text = para_text.replace("\n", '')  # remove empty rows
    return para_text


def thread_start(threads):
    for t in threads:
        t.setDaemon(True)
        t.start()
        sleep(2)
        # This case origin from Baidu Translator interface has the access controller
        # which does not allow the users to query the translator more frequently.
        # The sleep() is used to avoid the access limit error
    for t in threads:
        t.join()


def task(file_obj):
    res, result_list, translated_res = '', [], ''
    (filepath, tempfilename) = os.path.split(file_obj)
    filename, extension = os.path.splitext(tempfilename)
    wpath = filepath + '/' + filename + '_result' + extension

    if re.match(r'.+(docx)$', file_obj, re.M) and not re.match(r'.+\~\$.+', file_obj, re.M):  # docx format
        # exclude the case C:/My Received Files\~$test.docx
        print("task of " + file_obj + " start")
        doc_obj = docx.Document(file_obj)
        result_list = map(remove_doc_space, doc_obj.paragraphs)

        for para in result_list:
            if (para is not ''):  # remove empty rows
                res = res + para

        translated_res = BaiduTranslator.translate(res, fromLang='zh', toLang='en')
        wstr = translated_res['trans_result'][0]['dst'].encode("utf-8")
        document = Document()
        document.add_paragraph(wstr)
        document.save(wpath)

    if re.match(r'.+(txt)$', file_obj, re.M) and not re.match(r'.+\~\$.+', file_obj, re.M):  # txt format
        print("task of " + file_obj + " start")
        f = open(file_obj, "r")
        result_list = f.readlines()
        result_list = map(remove_txt_space, result_list)
        for para in result_list:
            if (para is not ''):  # remove empty rows
                res = res + para

        translated_res = BaiduTranslator.translate(res, fromLang='zh', toLang='en')
        wstr = translated_res['trans_result'][0]['dst'].encode("utf-8")
        wfile_obj = open(wpath, 'w')
        wfile_obj.write(wstr)
        f.close()
        wfile_obj.close()


def translate_from_path(path):
    re_str = "/*"
    glob_path = path + re_str
    files = glob.glob(pathname=glob_path)
    threads = []
    threads_num = 0
    for file in files:
        threads_num = threads_num + 1
        t = threading.Thread(target=task, args=(file,))
        threads.append(t)
        if threads_num == 10:  # 10 tasks are putting into one bulk for multi-processing
            threads_num = 0
            thread_start(threads)
            threads = []
        if file is files[-1]:
            thread_start(threads)

    # summarize tasks
    files_after = glob.glob(pathname=glob_path)
    wfile_obj = open(path + '/summary.txt', 'w+')
    restful_output = "Total tasks is " + str(len(files)) + " \n"

    for file in files:
        (filepath, tempfilename) = os.path.split(file)
        filename, extension = os.path.splitext(tempfilename)
        wpath = filepath + '\\' + filename + '_result' + extension;
        if wpath in files_after:
            wfile_obj.writelines(filename + extension + ": success  \n")
            restful_output = "{0}{1}{2}: success\n".format(restful_output, filename, extension)
        else:
            wfile_obj.writelines(filename + extension + ": fail \n")
            restful_output = "{0}{1}{2}: fail\n".format(restful_output, filename, extension)
    wfile_obj.close()
    return restful_output

    # translate_path('C:/Users/eqsvimp/PycharmProjects/Translator/testfiles')
    # print
